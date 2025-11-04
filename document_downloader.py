#!/usr/bin/env python3
"""
SAM.gov Belge İndirici ve Analiz Edici
"""

import requests
import os
import tempfile
import mimetypes
from urllib.parse import urlparse
import PyPDF2
import docx
from docx import Document as DocxDocument
import json
import re
from datetime import datetime

class DocumentDownloader:
    """SAM.gov belgelerini indiren ve analiz eden sınıf"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        self.temp_dir = tempfile.mkdtemp()
    
    def download_document(self, url, opportunity_id):
        """Belgeyi indir ve analiz et"""
        try:
            print(f"[DOWNLOAD] Belge indiriliyor: {url}")
            
            # URL'yi kontrol et
            if not url or not url.startswith(('http', 'file')):
                return None
            
            # File URL'leri için özel işleme
            if url.startswith('file://'):
                file_path = url[7:]  # file:// kısmını kaldır
                if not os.path.exists(file_path):
                    print(f"[ERROR] Dosya bulunamadı: {file_path}")
                    return None
                
                # Dosyayı oku
                with open(file_path, 'rb') as f:
                    content = f.read()
                
                # Content type'ı belirle
                file_extension = os.path.splitext(file_path)[1].lower()
                content_type = self._get_content_type_from_extension(file_extension)
                
                # Mock response objesi oluştur
                class MockResponse:
                    def __init__(self, content, content_type):
                        self.content = content
                        self.headers = {'content-type': content_type}
                
                response = MockResponse(content, content_type)
            else:
                # HTTP URL'leri için normal indirme
                response = self.session.get(url, timeout=30)
                response.raise_for_status()
            
            # Dosya uzantısını belirle
            content_type = response.headers.get('content-type', '')
            file_extension = self._get_file_extension(content_type, url)
            
            # Geçici dosya oluştur
            filename = f"{opportunity_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{file_extension}"
            filepath = os.path.join(self.temp_dir, filename)
            
            with open(filepath, 'wb') as f:
                f.write(response.content)
            
            print(f"[DOWNLOAD] Belge indirildi: {filename} ({len(response.content)} bytes)")
            
            # Belgeyi analiz et
            analysis = self._analyze_document(filepath, content_type)
            
            return {
                'filepath': filepath,
                'filename': filename,
                'content_type': content_type,
                'size': len(response.content),
                'analysis': analysis
            }
            
        except Exception as e:
            print(f"[ERROR] Belge indirme hatası: {e}")
            return None
    
    def _get_file_extension(self, content_type, url):
        """Dosya uzantısını belirle"""
        if 'pdf' in content_type.lower():
            return '.pdf'
        elif 'word' in content_type.lower() or 'document' in content_type.lower():
            return '.docx'
        elif 'excel' in content_type.lower() or 'spreadsheet' in content_type.lower():
            return '.xlsx'
        elif 'text' in content_type.lower():
            return '.txt'
        else:
            # URL'den uzantı çıkar
            parsed_url = urlparse(url)
            path = parsed_url.path.lower()
            if path.endswith('.pdf'):
                return '.pdf'
            elif path.endswith('.docx') or path.endswith('.doc'):
                return '.docx'
            elif path.endswith('.xlsx') or path.endswith('.xls'):
                return '.xlsx'
            elif path.endswith('.txt'):
                return '.txt'
            else:
                return '.bin'
    
    def _get_content_type_from_extension(self, file_extension):
        """Dosya uzantısından content type belirle"""
        extension_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.txt': 'text/plain'
        }
        return extension_map.get(file_extension, 'application/octet-stream')
    
    def _analyze_document(self, filepath, content_type):
        """Belgeyi analiz et ve içeriği çıkar"""
        analysis = {
            'text_content': '',
            'word_count': 0,
            'pages': 0,
            'sections': [],
            'tables': [],
            'images': 0,
            'metadata': {}
        }
        
        try:
            if content_type.lower().startswith('application/pdf'):
                analysis = self._analyze_pdf(filepath)
            elif 'word' in content_type.lower() or filepath.endswith('.docx'):
                analysis = self._analyze_docx(filepath)
            elif 'excel' in content_type.lower() or filepath.endswith('.xlsx'):
                analysis = self._analyze_excel(filepath)
            elif 'text' in content_type.lower() or filepath.endswith('.txt'):
                analysis = self._analyze_text(filepath)
            else:
                analysis['text_content'] = f"Desteklenmeyen dosya formatı: {content_type}"
                
        except Exception as e:
            analysis['text_content'] = f"Belge analiz hatası: {e}"
        
        return analysis
    
    def _analyze_pdf(self, filepath):
        """PDF belgesini analiz et"""
        analysis = {
            'text_content': '',
            'word_count': 0,
            'pages': 0,
            'sections': [],
            'tables': [],
            'images': 0,
            'metadata': {}
        }
        
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                analysis['pages'] = len(pdf_reader.pages)
                
                # Tüm sayfaları oku
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        full_text += page_text + "\n"
                    except Exception as e:
                        print(f"[WARNING] Sayfa {page_num + 1} okunamadı: {e}")
                
                analysis['text_content'] = full_text
                analysis['word_count'] = len(full_text.split())
                
                # Metadata
                if pdf_reader.metadata:
                    analysis['metadata'] = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                    }
                
                # Bölümleri tespit et
                analysis['sections'] = self._extract_sections(full_text)
                
        except Exception as e:
            analysis['text_content'] = f"PDF analiz hatası: {e}"
        
        return analysis
    
    def _analyze_docx(self, filepath):
        """Word belgesini analiz et"""
        analysis = {
            'text_content': '',
            'word_count': 0,
            'pages': 0,
            'sections': [],
            'tables': [],
            'images': 0,
            'metadata': {}
        }
        
        try:
            doc = DocxDocument(filepath)
            
            # Tüm paragrafları oku
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            analysis['text_content'] = full_text
            analysis['word_count'] = len(full_text.split())
            
            # Tabloları say
            analysis['tables'] = len(doc.tables)
            
            # Resimleri say
            analysis['images'] = len(doc.inline_shapes)
            
            # Bölümleri tespit et
            analysis['sections'] = self._extract_sections(full_text)
            
        except Exception as e:
            analysis['text_content'] = f"DOCX analiz hatası: {e}"
        
        return analysis
    
    def _analyze_excel(self, filepath):
        """Excel belgesini analiz et"""
        analysis = {
            'text_content': '',
            'word_count': 0,
            'pages': 0,
            'sections': [],
            'tables': [],
            'images': 0,
            'metadata': {}
        }
        
        try:
            import pandas as pd
            
            # Excel dosyasını oku
            excel_file = pd.ExcelFile(filepath)
            analysis['pages'] = len(excel_file.sheet_names)
            
            # Tüm sayfaları oku
            full_text = ""
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                full_text += f"Sayfa: {sheet_name}\n"
                full_text += df.to_string() + "\n\n"
            
            analysis['text_content'] = full_text
            analysis['word_count'] = len(full_text.split())
            analysis['tables'] = len(excel_file.sheet_names)
            
        except Exception as e:
            analysis['text_content'] = f"Excel analiz hatası: {e}"
        
        return analysis
    
    def _analyze_text(self, filepath):
        """Metin belgesini analiz et"""
        analysis = {
            'text_content': '',
            'word_count': 0,
            'pages': 0,
            'sections': [],
            'tables': [],
            'images': 0,
            'metadata': {}
        }
        
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            analysis['text_content'] = content
            analysis['word_count'] = len(content.split())
            analysis['pages'] = len(content) // 2000 + 1  # Tahmini sayfa sayısı
            
            # Bölümleri tespit et
            analysis['sections'] = self._extract_sections(content)
            
        except Exception as e:
            analysis['text_content'] = f"Metin analiz hatası: {e}"
        
        return analysis
    
    def _extract_sections(self, text):
        """Metinden bölümleri çıkar"""
        sections = []
        
        # Başlık desenleri
        heading_patterns = [
            r'^\d+\.\s+[A-Z][^\\n]*$',  # 1. BAŞLIK
            r'^[A-Z][A-Z\s]+$',        # BAŞLIK
            r'^Section\s+\d+',         # Section 1
            r'^Chapter\s+\d+',         # Chapter 1
            r'^Part\s+\d+',            # Part 1
        ]
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if len(line) > 5:  # En az 5 karakter
                for pattern in heading_patterns:
                    if re.match(pattern, line, re.MULTILINE):
                        sections.append({
                            'line_number': i + 1,
                            'text': line,
                            'level': self._get_heading_level(line)
                        })
                        break
        
        return sections
    
    def _get_heading_level(self, heading):
        """Başlık seviyesini belirle"""
        if re.match(r'^\d+\.\d+', heading):
            return 2
        elif re.match(r'^\d+\.', heading):
            return 1
        elif heading.isupper():
            return 0
        else:
            return 1
    
    def cleanup(self):
        """Geçici dosyaları temizle"""
        try:
            import shutil
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            print(f"[WARNING] Geçici dosyalar temizlenemedi: {e}")

# Test fonksiyonu
def test_document_downloader():
    """Document downloader'ı test et"""
    downloader = DocumentDownloader()
    
    # Test URL'leri
    test_urls = [
        "https://api.sam.gov/prod/opportunities/v1/noticedesc?noticeid=eea91e1b43be4ba4a24aa658c678a478",
        "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf"
    ]
    
    for url in test_urls:
        print(f"\n=== Test URL: {url} ===")
        result = downloader.download_document(url, "test_123")
        if result:
            print(f"Başarılı: {result['filename']}")
            print(f"Boyut: {result['size']} bytes")
            print(f"İçerik uzunluğu: {len(result['analysis']['text_content'])} karakter")
            print(f"Kelime sayısı: {result['analysis']['word_count']}")
        else:
            print("Başarısız")
    
    downloader.cleanup()

if __name__ == "__main__":
    test_document_downloader()
